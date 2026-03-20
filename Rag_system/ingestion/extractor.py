"""
Text extraction utilities.

Supports: PDF, DOCX, DOC, TXT, MD, images (OCR via pytesseract if available).
Returns plain text ready for LLM ingestion.
"""

from __future__ import annotations
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Max chars returned (keep within LLM context window)
MAX_CHARS = 50_000


def extract_text(file_path: str | Path) -> str:
    """
    Extract plain text from a file. Auto-detects format by extension.
    Returns empty string on failure (logs error).
    """
    path = Path(file_path)
    if not path.exists():
        logger.error(f"File not found: {file_path}")
        return ""

    suffix = path.suffix.lower()

    extractors = {
        ".pdf":  _from_pdf,
        ".docx": _from_docx,
        ".doc":  _from_docx,
        ".txt":  _from_text,
        ".md":   _from_text,
        ".png":  _from_image,
        ".jpg":  _from_image,
        ".jpeg": _from_image,
        ".tiff": _from_image,
        ".bmp":  _from_image,
    }

    extractor = extractors.get(suffix, _from_text)
    try:
        text = extractor(path)
        return text[:MAX_CHARS]
    except Exception as e:
        logger.error(f"Extraction failed for {file_path}: {e}")
        return ""


# ─── Format-specific extractors ───────────────────────────────────────────────

def _from_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"[Page {i+1}]\n{text}")
    result = "\n\n".join(pages)

    # If PDF has very little text it might be scanned — try OCR
    if len(result.strip()) < 100:
        logger.info(f"PDF appears scanned, attempting OCR: {path}")
        result = _pdf_ocr(path) or result

    return result


def _pdf_ocr(path: Path) -> str:
    """OCR a scanned PDF by converting pages to images first."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        images = convert_from_path(str(path), dpi=200)
        pages = [pytesseract.image_to_string(img) for img in images]
        return "\n\n".join(pages)
    except ImportError:
        logger.warning("pytesseract/pdf2image not installed; skipping OCR")
        return ""
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


def _from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _from_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _from_image(path: Path) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        return pytesseract.image_to_string(img)
    except ImportError:
        logger.warning("pytesseract/Pillow not installed; cannot OCR image")
        return ""
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return ""


# ─── Chunker (used by RAG ingestion) ─────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 800,
    overlap: int = 100,
) -> list[str]:
    """
    Split text into overlapping chunks for RAG ingestion.
    Tries to split on sentence/newline boundaries.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))

        # Try to find a natural break point (newline or period) near end
        if end < len(text):
            for sep in ["\n\n", "\n", ". ", " "]:
                idx = text.rfind(sep, start + chunk_size // 2, end)
                if idx != -1:
                    end = idx + len(sep)
                    break

        chunks.append(text[start:end])
        start = end - overlap

    return [c for c in chunks if c.strip()]
